#!/usr/bin/env python3

import os
from docx import Document
import markdown
import sys
import re
print(sys.executable)

# docx_path=sys.argv[1]  # Name of input file
docx_path = "LIS-Disaster Data Catalog Page.docx" # Example filename for testing

def empty_string_clause()


def extract_text_from_docx(docx_path):
    # Read the Word document
    doc = Document(docx_path)
    content = []
    
    # # Loop through the paragraphs and extract text
    # for para in doc.paragraphs:
    #     # For headings (h1, h2, h3), we will convert them to markdown headings
    #     if para.style.name.startswith('Heading'):
    #         heading_level = int(para.style.name.split()[-1])
    #         content.append(f"{'#' * heading_level} {para.text}\n")
    #     else:
    #         content.append(para.text + "\n")
    
    # Process lists (bulleted or numbered)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                content.append(cell.text.strip() + "\n")
    
    return "".join(content)


def parse_media_url(vals):
    '''Ensure that the media link is correctly formatted 
    and outputs are always saved to media directory'''
    if 'media' in vals:
        pattern = r'(\./media/[^\s]+)'
        match = re.search(pattern, vals)
        if match:
            vals = match.group(1)
        else:
            vals = f"./media/{vals.split(': ')[-1]}"

    return vals

def parse_media_alt_text(all_text):
    #Return additional information including alt text, author, and original url
    output = []
    match_alt_text = re.search(r'Image text \(alt\): (.*?)\n', all_text)
    match_author_name = re.search(r'Author name: (.*?)\n',all_text)
    match_author_URL = re.search(r'Author URL:\s*(.*)', all_text)

    out_names = ['main_fig_alt_text', 'main_fig_author_name', 'main_fig_author_URL']
    out_data = [match_alt_text, match_author_name, match_author_URL]

    for idx,match_ in enumerate(zip(out_names, out_data)):
        if match_[1]:
            output.append({match_[0]: match_[1].group(1)})
        else:
            output.append({match_[0]: "Data not provided"})
    return output


def parse_tag_information(all_text):
    #Return additional information including topics, data source, and match_prod_type
    output = []
    match_topic_values = re.search(r'Topic: (.*?)\n', all_text)
    match_source = re.search(r'Source: (.*?)\n',all_text)
    match_prod_type = re.search(r'Product Type:\s*(.*)', all_text)

    out_names = ['topic', 'source', 'product_type']
    out_data = [match_topic_values, match_source, match_prod_type]

    for idx,match_ in enumerate(zip(out_names, out_data)):
        if match_[1]:
            output.append({match_[0]: match_[1].group(1)})
        else:
            output.append({match_[0]: "Data not provided"})
    return output

#test 
# all_text = row.cells[1].text.strip()

def parse_layer_information(all_text):
    print(f'Parsing layer information. Currently code is only set for hex values.')
    #Return additional information including alt text, author, and original url
    match_layer_names = re.findall(r'Layer name: (.*?)\n', all_text)
    match_stacCol = re.findall(r'stacCol: (.*?)\n',all_text)
    match_layer_ids = re.findall(r'Layer id: (.*?)\n', all_text)
    match_layer_description = re.findall(r'Layer description: (.*?)\n', all_text)
    match_units = re.findall(r'Units: (.*?)\n', all_text)
    match_color_ramp_description = re.findall(r'Color ramp: (.*?)\n', all_text)
    match_color_stops = re.findall(r'\[([^\]]+)\]', all_text)

    # Clean up the groups (remove extra spaces and stray commas)
    final_color_groups = []
    for match in match_color_stops:
        items = [item.strip(" '\n") for item in match.split(",")]
        # Remove empty items or stray commas
        items = [item for item in items if item]
        final_color_groups.append(items)

    num_layers = len(match_layer_names)

    out_names = ['layer_name', 'stacCol', 'layer_id', 'layer_description', 'units', 'color_ramp_description', 'color_stops']
    out_data = [match_layer_names, match_stacCol, match_layer_ids,match_layer_description, match_units, match_color_ramp_description, final_color_groups]

    output = []
    for i in range(num_layers):
        layer_data = {}  # Dictionary for this specific layer
        for idx, match_ in enumerate(zip(out_names, out_data)):
            layer_data[f'{match_[0]}{i}'] = match_[1][i]
        output.append(layer_data)

    return output


def check_value_string_length(vals):
    #If string is empty, then replace the value with 
    return vals if len(vals) >=1 else "Data not provided"


def extract_table_info_from_docx(docx_path):
    # Read the Word document
    doc = Document(docx_path)
    extracted_data = {}
    
    # Process lists (bulleted or numbered)
    for table in doc.tables:
        for iRow,row in enumerate(table.rows):
            header = row.cells[0].text.strip().lower().split("\n")[0]
            assert len(header) != 0, "Header is empty, need to check the template document and update it to include all headers"
            print(f'Index {iRow} is header {header}')
            if iRow == 5:
                break

            if header == "media":
                all_text = parse_media_alt_text(row.cells[1].text.strip())
                vals = parse_media_url(row.cells[1].text.strip().split("\n")[0])
                
                if header not in extracted_data:
                    extracted_data[header] = []

                extracted_data[header].append({'main_media_image': check_value_string_length(vals)})
                for idx, val in enumerate(all_text):
                    extracted_data[header].append({next(iter(val.keys())):val[next(iter(val.keys()))]})

            elif header == 'tags':
                all_text = parse_tag_information(row.cells[1].text.strip())
                if header not in extracted_data:
                    extracted_data[header] = []

                for idx, val in enumerate(all_text):
                    extracted_data[header].append({next(iter(val.keys())):val[next(iter(val.keys()))]})
            elif header == 'layers':
                all_text = parse_layer_information(row.cells[1].text.strip())

                for i in range(len(all_text)):
                        
                    if header not in extracted_data:
                        extracted_data[f'{header}'] = []

                    extracted_data[f'{header}'].append({f'Layer{i}':all_text[i]})


            else:
                vals = row.cells[1].text.strip().split("\n")[0]
                extracted_data[header] = check_value_string_length(vals)

    return "".join(content)

def convert_to_mdx(docx_path, mdx_path):
    # Extract the text from the DOCX
    content = extract_text_from_docx(docx_path)
    content_table = extract_table_info_from_docx(docx_path)

    # Optionally: Convert extracted content to markdown format
    # (You could use `markdown` module here if necessary for extra conversions)
    md_content = markdown.markdown(content)

    # Prepare the .mdx content (in case you need custom JSX formatting)
    mdx_content = f"import MyComponent from './MyComponent'\n\n{md_content}"

    # Save the MDX file
    with open('test.mdx', "w") as mdx_file:
        mdx_file.write(mdx_content)

    print(f"MDX file created at: {mdx_path}")

# Example usage
mdx_file = "output.mdx"
convert_to_mdx(filename, mdx_file)
