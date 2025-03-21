#!/usr/bin/env python3

import os
from docx import Document
import markdown
import sys
import re
import yaml
print(sys.executable)


def read_document(docx_path):
    """Read a Word document with error handling."""
    # Check if the file exists
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"Error: File '{docx_path}' not found.")

    # Check if the file is a .docx file
    if not docx_path.lower().endswith('.docx'):
        raise ValueError(f"Error: File '{docx_path}' is not a .docx file.")

    # Check if the file is not empty
    if os.path.getsize(docx_path) == 0:
        raise ValueError(f"Error: File '{docx_path}' is empty.")

    try:
        print(f'Opening {docx_path}')
        doc = Document(docx_path)
        print(f"Successfully opened '{docx_path}'.")
        return doc
    except Exception as e:
        raise IOError(f"Error reading file '{docx_path}': {e}")



def parse_media_url(vals):
    '''Ensure that the media link is correctly formatted 
    and outputs are always saved to media directory'''
    if 'media' in vals:
        pattern = r'(\./media/[^\s]+)'
        match = re.search(pattern, vals)
        if match:
            vals = match.group(1)
        else:
            vals = f"./media/{vals.split(': ')[-1]}"

    return vals

def parse_media_alt_text(all_text):
    #Return additional information including alt text, author, and original url
    output = []
    match_alt_text = re.search(r'Image text \(alt\): (.*?)\n', all_text, re.IGNORECASE)
    match_author_name = re.search(r'Author name: (.*?)\n',all_text, re.IGNORECASE)
    match_author_URL = re.search(r'Author URL:\s*(.*)', all_text, re.IGNORECASE)

    out_names = ['main_fig_alt_text', 'main_fig_author_name', 'main_fig_author_URL']
    out_data = [match_alt_text, match_author_name, match_author_URL]

    for idx,match_ in enumerate(zip(out_names, out_data)):
        if match_[1]:
            output.append({match_[0]: match_[1].group(1)})
        else:
            output.append({match_[0]: "Data not provided"})
    return output


def parse_tag_information(all_text):
    #Return additional information including topics, data source, and match_prod_type
    output = []
    match_topic_values = re.search(r'Topic: (.*?)\n', all_text, re.IGNORECASE)
    match_source = re.search(r'Source: (.*?)\n',all_text, re.IGNORECASE)
    match_prod_type = re.search(r'Product Type:\s*(.*)', all_text, re.IGNORECASE)

    out_names = ['topic', 'source', 'product_type']
    out_data = [match_topic_values, match_source, match_prod_type]

    for idx,match_ in enumerate(zip(out_names, out_data)):
        if match_[1]:
            output.append({match_[0]: match_[1].group(1)})
        else:
            output.append({match_[0]: "Data not provided"})
    return output


def parse_layer_information(all_text):
    #Return additional information including alt text, author, and original url
    layer_name = re.findall(r'Layer name: (.*?)\n', all_text, re.IGNORECASE)
    stacCol = re.findall(r'stacCol: (.*?)\n',all_text, re.IGNORECASE)
    layer_id = re.findall(r'Layer id: (.*?)\n', all_text, re.IGNORECASE)
    layer_description = re.findall(r'Layer description: (.*?)\n', all_text, re.IGNORECASE)
    unit = re.findall(r'Units: (.*?)\n', all_text, re.IGNORECASE)
    color_ramp_description = re.findall(r'Color ramp: (.*?)\n', all_text, re.IGNORECASE)
    color_stops = re.findall(r'\[([^\]]+)\]', all_text, re.IGNORECASE)
    data_format = re.findall(r'Data format: (.*?)\n', all_text, re.IGNORECASE)
    projection = re.findall(r'Projection: (.*?)\n', all_text, re.IGNORECASE)
    legend_min = re.findall(r'Legend minimum: (.*?)\n', all_text, re.IGNORECASE)
    legend_max = re.findall(r'Legend maximum: (.*?)\n', all_text, re.IGNORECASE)
    legend_type = re.findall(r'Legend type: (.*?)\n', all_text, re.IGNORECASE)

    # Clean up the groups (remove extra spaces and stray commas)
    final_color_groups = []
    for match in color_stops:
        items = [item.strip(" '\n") for item in match.split(",")]
        # Remove empty items or stray commas
        items = [item for item in items if item]
        final_color_groups.append(items)

    num_layers = len(layer_name)

    #This can allow for the selection of specific information into the .mdx file without adding everything
    out_names = ['layer_name', 'stacCol', 'layer_id', 'layer_description', 'units', 'color_ramp_description', 'color_stops', 'data_format','projection','legend_min','legend_max','legend_type']
    out_data = [layer_name, stacCol, layer_id, layer_description, unit, color_ramp_description, final_color_groups, data_format, projection, legend_min, legend_max, legend_type]

    output = []
    for i in range(num_layers):
        layer_data = {}  # Dictionary for this specific layer
        for idx, match_ in enumerate(zip(out_names, out_data)):
            try:
                layer_data[f'{match_[0]}{i}'] = match_[1][i]
            except IndexError:
                pass
        output.append(layer_data)

    return output


def table_0_info(row, header, extracted_data):
    #Extract information from the first table of the document.
    if header == "media":
        all_text = parse_media_alt_text(row.cells[1].text.strip())
        vals = parse_media_url(row.cells[1].text.strip().split("\n")[0])
        
        if header not in extracted_data:
            extracted_data[header] = []

        extracted_data[header].append({'main_media_image': check_value_string_length(vals)})
        for idx, val in enumerate(all_text):
            extracted_data[header].append({next(iter(val.keys())):val[next(iter(val.keys()))]})

    elif header == 'tags':
        all_text = parse_tag_information(row.cells[1].text.strip())
        if header not in extracted_data:
            extracted_data[header] = []

        for idx, val in enumerate(all_text):
            extracted_data[header].append({next(iter(val.keys())):val[next(iter(val.keys()))]})
    elif header == 'layers':
        print(f'Parsing layer information.')
        all_text = parse_layer_information(row.cells[1].text.strip())

        for i in range(len(all_text)):
                
            if header not in extracted_data:
                extracted_data[f'{header}'] = []

            extracted_data[f'{header}'].append({f'Layer{i}':all_text[i]})
    else:
        vals = row.cells[1].text.strip().split("\n")[0]
        extracted_data[header] = check_value_string_length(vals)

    return extracted_data

def check_value_string_length(vals):
    #If string is empty, then replace the value with 
    return vals if len(vals) >=1 else "Data not provided"

def parse_table_value_content(row,header,table_0,table_1):
    all_text = check_value_string_length(row.cells[1].text.strip())
    
    #Returns a list of items after Value:
    if header == 'content_source':
        if re.findall(r'(?<=Value:\s)(.*)', all_text, re.IGNORECASE)[0] == 'null':
            source_values = [item['source'] for item in table_0['tags'] if 'source' in item]
            table_1[header] = source_values
        else:
            match_value_names = re.findall(r'(?<=Value:\s)(.*)', all_text, re.IGNORECASE)
            table_1[header] = match_value_names
    elif header == 'temporal_extent':
        match_value_names = re.findall(r'Start:\s*(\d{2}/\d{2}/\d{4})|End:\s*(\d{2}/\d{2}/\d{4})', all_text, re.IGNORECASE)
        # Extract values
        start_value = next((match[0] for match in match_value_names if match[0]), None)
        end_value = next((match[1] for match in match_value_names if match[1]), None)
        table_1['start_temporal_extent'] = start_value
        table_1['end_temporal_extent'] =end_value
    elif header == 'legend_value_range':
         pass
        #This will add if all of the layers are identical, but that isn't always the case
        #  match_value_names = re.findall(r"Min:\s*([\d.]+)\s*Max:\s*([\d.]+)\s*Type:\s*(\w+)", all_text)
        #  min_value = next((match[0] for match in match_value_names if match[0]), None)
        #  max_value = next((match[1] for match in match_value_names if match[1]), None)
        #  legend_type = next((match[2] for match in match_value_names if match[2]), None)
        #  table_1['legend_min'] = min_value
        #  table_1['legend_max'] = max_value
        #  table_1['legend_type'] = legend_type
    else:

        match_value_names = re.findall(r'(?<=Value:\s)(.*)', all_text, re.IGNORECASE)
        table_1[header] = match_value_names
    return table_1

def parse_additional_table_info(row, header, table_2):

    all_text = check_value_string_length(row.cells[1].text.strip())
    # Regex pattern to capture header and value separately
    match = re.search(r'Header:\s*(.*?)\s*\n+\s*Value:\s*(.*)', all_text, re.DOTALL)
    
    if match:
        header_ = match.group(1).strip()
        value = match.group(2).strip()
        if len(header_) != 0 and len(value) != 0:
            table_2[header] = []
            table_2[header].append({header_: value})
    return table_2


def extract_table_info_from_docx(doc):
    #Extract information from the tables in the Word document.
    table_0 = {}
    table_1 = {} #prose block with general dataset information
    table_2 = {} #prose blocks with additional information
    
    # Process lists (bulleted or numbered)
    for iTable,table in enumerate(doc.tables):
        for iRow,row in enumerate(table.rows):
            header = row.cells[0].text.strip().lower().split("\n")[0]

            # Skip empty rows and headers
            if len(header) == 0 and len(row.cells[1].text.strip()) == 0:
                continue
            elif len(header) == 0:
                assert len(header) != 0, "Header is empty, need to check the template document and update it to include all headers"
            else:
                if iTable == 0:
                    table_0 = table_0_info(row, header, table_0)
                elif iTable == 1:
                    table_1 = parse_table_value_content(row,header,table_0,table_1)
                elif iTable ==2:
                    table_2 = parse_additional_table_info(row, header, table_2)
    return table_0, table_1, table_2


def extract_headers_and_paragraphs(doc):
    output = {}
    current_header = None
    current_content = []

    for para in doc.paragraphs:
        # Detect headers based on style or bold formatting
        if para.style.name.startswith('Heading') or any(run.bold for run in para.runs):
            if current_header:
                # Save previous header-content pair to the dictionary
                output[current_header] = "\n".join(current_content).strip()
                current_content = []  # Reset content list
            current_header = para.text.strip()
        elif current_header:
            current_content.append(para.text.strip())

    # Save the last header-content pair
    if current_header and current_content:
        output[current_header] = "\n".join(current_content).strip()

    return output


def retrieve_all_docx_data(docx_path):
    try:
        doc = read_document(docx_path)
    except Exception as e:
        print(e)
    # Extract the text from the DOCX
    table_0, table_1, table_2 = extract_table_info_from_docx(doc)
    content = extract_headers_and_paragraphs(doc)

    return table_0, table_1, table_2, content



